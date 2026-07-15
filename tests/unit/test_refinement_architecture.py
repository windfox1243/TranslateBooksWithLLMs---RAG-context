from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest
from flask import Flask

from src.core.llm.base import LLMResponse
from src.persistence.checkpoint_manager import CheckpointManager
from src.utils.novel_context import (
    RefinementContextTracker,
    build_novel_context,
    compress_dynamic_state,
    map_context_snapshots_for_refinement,
    normalize_refinement_context,
)


@pytest.mark.asyncio
async def test_docx_plain_checkpoint_rebuild_uses_effective_chunk_text(tmp_path):
    from io import BytesIO
    from docx import Document
    from src.core.adapters import build_translated_output

    source_path = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Hello")
    document.add_paragraph("World")
    document.save(source_path)

    manager = CheckpointManager(str(tmp_path / "jobs.db"))
    assert manager.db.create_job("job", "docx", {
        "preserved_input_path": str(source_path),
        "output_file_path": str(tmp_path / "translated.docx"),
        "max_tokens_per_chunk": 450,
        "prompt_options": {"plain_text_mode": True},
    })
    assert manager.save_checkpoint(
        translation_id="job",
        chunk_index=0,
        original_text="Hello\n\nWorld",
        translated_text="Bonjour\n\nMonde",
        chunk_status="completed",
    )

    output, error = await build_translated_output("job", manager)

    assert error is None
    rebuilt = Document(BytesIO(output))
    assert [paragraph.text for paragraph in rebuilt.paragraphs] == [
        "Bonjour", "Monde"
    ]


def test_full_context_snapshots_map_without_becoming_nested_dynamic_state():
    full_context = build_novel_context("GLOBAL LORE", "HISTORICAL STATE")
    db_chunks = [
        {
            "chunk_index": 0,
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(full_context),
            },
        }
    ]

    mapped = map_context_snapshots_for_refinement(
        total_chunks=2,
        db_chunks=db_chunks,
        fallback_context=build_novel_context("FALLBACK", ""),
    )

    assert mapped == [full_context, full_context]
    assert mapped[0].count("---DYNAMIC_STATE_START---") == 1


def test_refinement_context_mapping_uses_translated_text_position():
    context_a = build_novel_context("CONTEXT A", "STATE A")
    context_b = build_novel_context("CONTEXT B", "STATE B")
    db_chunks = [
        {
            "chunk_index": 0,
            "translated_text": "a" * 100,
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(context_a),
            },
        },
        {
            "chunk_index": 1,
            "translated_text": "b" * 10,
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(context_b),
            },
        },
    ]

    mapped = map_context_snapshots_for_refinement(
        total_chunks=3,
        db_chunks=db_chunks,
        fallback_context="",
        refinement_units=["x" * 50, "y" * 50, "z" * 10],
    )

    assert mapped == [context_a, context_a, context_b]


def test_refinement_context_mapping_ignores_failed_or_partial_snapshots():
    context_a = build_novel_context("CONTEXT A", "STATE A")
    failed_context = build_novel_context("FAILED", "POISONED STATE")
    partial_context = build_novel_context("PARTIAL", "PARTIAL STATE")
    context_b = build_novel_context("CONTEXT B", "STATE B")
    db_chunks = [
        {
            "chunk_index": 0,
            "status": "completed",
            "translated_text": "translated a",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(context_a),
            },
        },
        {
            "chunk_index": 1,
            "status": "failed",
            "translated_text": "source fallback",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(failed_context),
            },
        },
        {
            "chunk_index": 2,
            "status": "partial",
            "translated_text": "source fallback",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(partial_context),
            },
        },
        {
            "chunk_index": 3,
            "status": "completed",
            "translated_text": "translated b",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(context_b),
            },
        },
    ]

    mapped = map_context_snapshots_for_refinement(
        total_chunks=2,
        db_chunks=db_chunks,
        fallback_context="",
    )

    assert mapped == [context_a, context_b]


def test_refinement_context_uses_final_lore_with_historical_dynamic_state():
    historical = build_novel_context(
        (
            "# GLOBAL LORE\n\n"
            "## CHARACTERS & GENDERS\n"
            "- Kriha: Unspecified, Captain.\n\n"
            "## GLOSSARY & TERMINOLOGY\n"
        ),
        (
            "## CURRENT ADDRESSING FORMS\n"
            "- Kriha → Valentine: source form \"Major\" | formal\n\n"
            "## RELATIONSHIP EVOLUTION\n"
            "- Kriha → Valentine: Newly assigned subordinate."
        ),
    )
    final_context = build_novel_context(
        (
            "# GLOBAL LORE\n\n"
            "## CHARACTERS & GENDERS\n"
            "- Kriha: Female, Captain and loyal subordinate.\n\n"
            "## GLOSSARY & TERMINOLOGY\n"
            "- blood art: huyết thuật\n"
        ),
        (
            "## CURRENT ADDRESSING FORMS\n"
            "- Kriha → Valentine: source form \"Valentine\" | intimate\n\n"
            "## RELATIONSHIP EVOLUTION\n"
            "- Kriha → Valentine: Deep mutual trust."
        ),
    )

    combined = normalize_refinement_context(historical, final_context)

    assert "- Kriha: Female, Captain and loyal subordinate." in combined
    assert "- blood art: huyết thuật" in combined
    assert 'source form "Major" | formal' in combined
    assert "Newly assigned subordinate" in combined
    assert 'source form "Valentine" | intimate' not in combined
    assert "Deep mutual trust" not in combined
    assert "Unspecified" not in combined


def test_context_editor_scope_controls_lore_leakage(
    monkeypatch,
    tmp_path,
):
    from src.api.blueprints.translation_routes import (
        create_translation_blueprint,
    )
    from src.utils.novel_context import save_novel_context
    import src.config

    historical = build_novel_context(
        (
            "# GLOBAL LORE\n\n"
            "## CHARACTERS & GENDERS\n"
            "- Kriha: Unspecified, Captain.\n\n"
            "## GLOSSARY & TERMINOLOGY\n"
        ),
        (
            "## CURRENT ADDRESSING FORMS\n"
            "- Kriha → Valentine: source form \"Major\" | formal"
        ),
    )
    latest = build_novel_context(
        (
            "# GLOBAL LORE\n\n"
            "## CHARACTERS & GENDERS\n"
            "- Kriha: Female, Captain and loyal subordinate.\n\n"
            "## GLOSSARY & TERMINOLOGY\n"
            "- blood art: huyết thuật\n"
        ),
        (
            "## CURRENT ADDRESSING FORMS\n"
            "- Kriha → Valentine: source form \"Valentine\" | intimate"
        ),
    )
    save_novel_context("novel.txt", tmp_path, latest)
    monkeypatch.setattr(src.config, "NOVEL_CONTEXTS_DIR", tmp_path)

    checkpoint_manager = MagicMock()
    checkpoint_manager.load_checkpoint.return_value = {
        "job": {
            "config": {
                "prompt_options": {
                    "novel_context_file": "novel.txt",
                }
            }
        },
        "chunks": [{
            "chunk_index": 0,
            "status": "completed",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(historical),
            },
        }],
    }
    state_manager = MagicMock()
    state_manager.checkpoint_manager = checkpoint_manager

    app = Flask(__name__)
    app.register_blueprint(
        create_translation_blueprint(
            state_manager,
            lambda *_args, **_kwargs: None,
            str(tmp_path),
        )
    )

    with app.test_client() as client:
        snapshot_response = client.get("/api/translation/job/context/0")
        global_response = client.get(
            "/api/translation/job/context/0?scope=global_lore"
        )

    assert snapshot_response.status_code == 200
    snapshot_content = snapshot_response.get_json()["context_content"]
    assert "- Kriha: Unspecified, Captain." in snapshot_content
    assert "- blood art: huyết thuật" not in snapshot_content
    assert 'source form "Major" | formal' in snapshot_content

    assert global_response.status_code == 200
    global_content = global_response.get_json()["context_content"]
    assert "- Kriha: Female, Captain and loyal subordinate." in global_content
    assert "- blood art: huyết thuật" in global_content
    assert 'source form "Major" | formal' in global_content
    assert 'source form "Valentine" | intimate' not in global_content
    assert "Unspecified" not in global_content


def test_context_resync_route_accepts_numeric_context_revision(
    monkeypatch,
    tmp_path,
):
    from src.api.blueprints.translation_routes import (
        create_translation_blueprint,
    )

    edited = build_novel_context(
        "# GLOBAL LORE",
        "## RELATIONSHIP EVOLUTION\n- Alice ↔ Bob: Close friends.",
    )
    checkpoint_manager = MagicMock()
    checkpoint_manager.load_checkpoint.return_value = {
        "job": {
            "config": {
                "prompt_options": {},
                "refine_after": False,
            },
            "progress": {"status": "completed"},
        },
        "chunks": [{
            "chunk_index": 0,
            "status": "completed",
            "original_text": "Source",
            "translated_text": "Translation",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(edited),
            },
        }],
    }
    checkpoint_manager.mark_refinement_stale.return_value = 7
    checkpoint_manager.db = MagicMock()

    state_manager = MagicMock()
    state_manager.checkpoint_manager = checkpoint_manager
    state_manager.get_translation.return_value = None

    thread = MagicMock()
    monkeypatch.setattr(
        "src.api.blueprints.translation_routes._claim_context_resync",
        lambda _translation_id: True,
    )
    monkeypatch.setattr(
        "src.api.blueprints.translation_routes._release_context_resync",
        lambda _translation_id: None,
    )
    monkeypatch.setattr(
        "src.api.blueprints.translation_routes.threading.Thread",
        lambda **_kwargs: thread,
    )

    app = Flask(__name__)
    app.register_blueprint(
        create_translation_blueprint(
            state_manager,
            lambda *_args, **_kwargs: None,
            str(tmp_path),
        )
    )

    with app.test_client() as client:
        response = client.post(
            "/api/translation/job/context/0/resync",
            json={"context_content": edited},
        )

    assert response.status_code == 200
    assert response.get_json()["context_revision"] == 7
    checkpoint_manager.mark_refinement_stale.assert_called_once_with("job")
    thread.start.assert_called_once()


def test_context_resync_route_accepts_failed_source_snapshot(
    monkeypatch,
    tmp_path,
):
    from src.api.blueprints.translation_routes import (
        create_translation_blueprint,
    )

    edited = build_novel_context(
        "# GLOBAL LORE",
        "## RELATIONSHIP EVOLUTION\n- Alice ↔ Bob: Close friends.",
    )
    checkpoint_manager = MagicMock()
    checkpoint_manager.load_checkpoint.return_value = {
        "job": {
            "config": {
                "prompt_options": {},
                "refine_after": False,
            },
            "progress": {"status": "partial"},
        },
        "chunks": [{
            "chunk_index": 2,
            "status": "failed",
            "original_text": "Source with useful context",
            "translated_text": "Source with useful context",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(edited),
            },
        }],
    }
    checkpoint_manager.mark_refinement_stale.return_value = 3
    checkpoint_manager.db = MagicMock()

    state_manager = MagicMock()
    state_manager.checkpoint_manager = checkpoint_manager
    state_manager.get_translation.return_value = None

    thread = MagicMock()
    monkeypatch.setattr(
        "src.api.blueprints.translation_routes._claim_context_resync",
        lambda _translation_id: True,
    )
    monkeypatch.setattr(
        "src.api.blueprints.translation_routes._release_context_resync",
        lambda _translation_id: None,
    )
    monkeypatch.setattr(
        "src.api.blueprints.translation_routes.threading.Thread",
        lambda **_kwargs: thread,
    )

    app = Flask(__name__)
    app.register_blueprint(
        create_translation_blueprint(
            state_manager,
            lambda *_args, **_kwargs: None,
            str(tmp_path),
        )
    )

    with app.test_client() as client:
        available = client.get("/api/translation/job/context/2")
        response = client.post(
            "/api/translation/job/context/2/resync",
            json={"context_content": edited},
        )

    assert available.status_code == 200
    assert available.get_json()["available_chunk_indices"] == [2]
    assert response.status_code == 200
    assert response.get_json()["context_revision"] == 3
    checkpoint_manager.db.save_chunk.assert_called_once()
    thread.start.assert_called_once()


def test_resume_translation_blocks_unfinished_context_resync(tmp_path):
    from src.api.blueprints.translation_routes import (
        create_translation_blueprint,
    )

    checkpoint_manager = MagicMock()
    checkpoint_manager.load_checkpoint.return_value = {
        "job": {
            "config": {
                "file_path": str(tmp_path / "book.txt"),
                "preserved_input_path": str(tmp_path / "book.txt"),
                "_context_resync": {
                    "status": "paused",
                    "last_processed_chunk": 12,
                },
            },
        },
        "resume_from_index": 13,
    }
    state_manager = MagicMock()
    state_manager.checkpoint_manager = checkpoint_manager
    state_manager.get_all_translations.return_value = {}

    app = Flask(__name__)
    app.register_blueprint(
        create_translation_blueprint(
            state_manager,
            lambda *_args, **_kwargs: None,
            str(tmp_path),
        )
    )

    with app.test_client() as client:
        response = client.post("/api/resume/job")

    assert response.status_code == 409
    assert response.get_json()["context_resync_state"]["status"] == "paused"
    state_manager.restore_job_from_checkpoint.assert_not_called()
    checkpoint_manager.mark_running.assert_not_called()


def test_resumable_jobs_expose_stale_context_resync_as_paused(tmp_path):
    from src.api.blueprints.translation_routes import (
        create_translation_blueprint,
    )

    config = {
        "file_path": str(tmp_path / "book.txt"),
        "output_filename": "book.out.txt",
        "openai_api_key": "YOUR_API_KEY_HERE",
        "_context_resync": {
            "status": "running",
            "last_processed_chunk": 7,
        },
    }
    checkpoint_manager = MagicMock()
    checkpoint_manager.get_job.return_value = {"config": config}
    checkpoint_manager.update_job_config.return_value = True
    state_manager = MagicMock()
    state_manager.checkpoint_manager = checkpoint_manager
    state_manager.get_resumable_jobs.return_value = [{
        "translation_id": "job",
        "status": "interrupted",
        "file_type": "txt",
        "config": dict(config),
        "progress": {"completed_chunks": 7, "total_chunks": 10},
    }]

    app = Flask(__name__)
    app.register_blueprint(
        create_translation_blueprint(
            state_manager,
            lambda *_args, **_kwargs: None,
            str(tmp_path),
        )
    )

    with app.test_client() as client:
        response = client.get("/api/resumable")

    assert response.status_code == 200
    job = response.get_json()["resumable_jobs"][0]
    assert job["context_resync"]["status"] == "paused"
    assert "openai_api_key" not in job["config"]
    checkpoint_manager.update_job_config.assert_called_once()


def test_context_resync_resume_restores_auto_resume_follow_up(
    monkeypatch,
    tmp_path,
):
    from src.api.blueprints.translation_routes import (
        create_translation_blueprint,
    )

    checkpoint_manager = MagicMock()
    checkpoint_manager.load_checkpoint.return_value = {
        "job": {
            "config": {
                "llm_provider": "openai",
                "llm_api_endpoint": "http://localhost:1234/v1",
                "_context_resync": {
                    "status": "paused",
                    "last_processed_chunk": 2,
                    "follow_up_kind": "auto_resume_translation",
                    "mode": "global_lore",
                },
            },
        },
        "resume_from_index": 3,
        "chunks": [{
            "chunk_index": 2,
            "status": "completed",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state("Context"),
            },
        }],
    }
    checkpoint_manager.update_job_config.return_value = True
    state_manager = MagicMock()
    state_manager.checkpoint_manager = checkpoint_manager

    captured = {}

    def fake_resync(*args):
        captured["auto_resume_callback"] = args[5]

    class ImmediateThread:
        def __init__(self, target, **_kwargs):
            self.target = target
            self.daemon = False

        def start(self):
            self.target()

    monkeypatch.setattr(
        "src.api.blueprints.translation_routes._claim_context_resync",
        lambda _translation_id: True,
    )
    monkeypatch.setattr(
        "src.api.blueprints.translation_routes._release_context_resync",
        lambda _translation_id: None,
    )
    monkeypatch.setattr(
        "src.core.adapters.generic_translator.resync_context_snapshots_background",
        fake_resync,
    )
    monkeypatch.setattr(
        "src.api.blueprints.translation_routes.threading.Thread",
        ImmediateThread,
    )

    app = Flask(__name__)
    app.register_blueprint(
        create_translation_blueprint(
            state_manager,
            lambda *_args, **_kwargs: None,
            str(tmp_path),
        )
    )

    with app.test_client() as client:
        response = client.post("/api/translation/job/context/resync/resume")

    assert response.status_code == 200
    assert callable(captured["auto_resume_callback"])


def test_refinement_source_survives_completed_checkpoint_cleanup(tmp_path):
    manager = CheckpointManager(db_path=str(tmp_path / "jobs.db"))
    manager.uploads_dir = tmp_path / "uploads"
    manager.uploads_dir.mkdir()

    original = tmp_path / "novel.txt"
    original.write_text("Source novel", encoding="utf-8")
    translated = tmp_path / "translated.txt"
    translated.write_text("First-pass translation", encoding="utf-8")
    config = {
        "file_type": "txt",
        "file_path": str(original),
        "output_filename": translated.name,
        "refine_after": True,
    }

    assert manager.start_job(
        "job",
        "txt",
        config,
        str(original),
    )
    preserved_original = Path(config["preserved_input_path"])
    refinement_source = manager.preserve_refinement_source(
        "job",
        str(translated),
        config,
    )

    assert refinement_source is not None
    assert Path(refinement_source).read_text(
        encoding="utf-8"
    ) == "First-pass translation"
    assert manager.cleanup_completed_job("job")
    assert Path(refinement_source).is_file()
    assert not preserved_original.exists()

    persisted = manager.load_checkpoint("job")["job"]["config"]
    assert persisted["refinement_source_path"] == refinement_source
    assert persisted["refinement_stale"] is True

    assert manager.mark_refinement_current("job")
    persisted = manager.load_checkpoint("job")["job"]["config"]
    assert persisted["refinement_stale"] is False
    assert persisted["refinement_context_revision"] == 0

    assert manager.mark_refinement_stale("job") == 1
    persisted = manager.load_checkpoint("job")["job"]["config"]
    assert persisted["refinement_stale"] is True
    assert persisted["context_revision"] == 1


def test_context_resync_correction_replays_preserved_first_pass(tmp_path):
    from src.api.blueprints.translation_routes import (
        _build_corrective_refinement_config,
    )

    source = tmp_path / "first-pass.epub"
    source.write_bytes(b"draft")
    output = tmp_path / "final.epub"
    output.write_bytes(b"refined")
    config = {
        "refine_after": True,
        "refinement_source_path": str(source),
        "output_filepath": str(output),
        "output_filename": output.name,
        "refine_only": False,
    }

    correction = _build_corrective_refinement_config(config)

    assert correction is not None
    assert correction["file_path"] == str(source.resolve())
    assert correction["_force_output_filepath"] == str(output.resolve())
    assert correction["refine_only"] is True
    assert correction["refine_after"] is False
    assert correction["_context_resync_refinement"] is True
    assert config["refine_only"] is False


def test_context_resync_does_not_refine_an_already_refined_output(tmp_path):
    from src.api.blueprints.translation_routes import (
        _build_corrective_refinement_config,
    )

    output = tmp_path / "final.txt"
    output.write_text("Already refined", encoding="utf-8")

    assert _build_corrective_refinement_config({
        "refine_after": True,
        "output_filepath": str(output),
    }) is None


@pytest.mark.asyncio
async def test_early_refinement_unit_receives_final_lore_and_early_state():
    historical = build_novel_context(
        (
            "# GLOBAL LORE\n\n"
            "## CHARACTERS & GENDERS\n"
            "- Kriha: Unspecified, Captain.\n\n"
            "## GLOSSARY & TERMINOLOGY\n"
        ),
        (
            "## CURRENT ADDRESSING FORMS\n"
            "- Kriha → Valentine: source form \"Major\" | formal\n\n"
            "## RELATIONSHIP EVOLUTION\n"
            "- Kriha → Valentine: Newly assigned subordinate."
        ),
    )
    final_context = build_novel_context(
        (
            "# GLOBAL LORE\n\n"
            "## CHARACTERS & GENDERS\n"
            "- Kriha: Female, Captain and loyal subordinate.\n\n"
            "## GLOSSARY & TERMINOLOGY\n"
            "- blood art: huyết thuật\n"
        ),
        "",
    )
    tracker = RefinementContextTracker(
        prompt_options={"novel_context": final_context},
        historical_contexts=[historical],
    )

    combined = await tracker.next_context(
        text="Draft chapter one.",
        llm_client=MagicMock(),
        model_name="model",
        target_language="English",
        display_index=1,
        total_chunks=1,
        scene_key=0,
    )

    assert "- Kriha: Female, Captain and loyal subordinate." in combined
    assert "- blood art: huyết thuật" in combined
    assert 'source form "Major" | formal' in combined
    assert "Newly assigned subordinate" in combined


@pytest.mark.asyncio
async def test_legacy_refinement_snapshots_recover_omitted_dormant_relationships():
    first_snapshot = build_novel_context(
        "# GLOBAL LORE",
        (
            "## CURRENT ADDRESSING FORMS\n"
            '- Alice → Bob: source form "Bob" | target-language form "anh" | intimate\n\n'
            "## RELATIONSHIP EVOLUTION\n"
            "- Alice ↔ Bob: Established romantic couple."
        ),
    )
    legacy_later_snapshot = build_novel_context(
        "# GLOBAL LORE",
        (
            "## CURRENT ADDRESSING FORMS\n"
            '- Guard → Captain: source form "Captain" | formal\n\n'
            "## RELATIONSHIP EVOLUTION\n"
            "- Guard → Captain: Temporary scene relationship."
        ),
    )
    tracker = RefinementContextTracker(
        prompt_options={
            "novel_context": build_novel_context(
                "# GLOBAL LORE",
                "## RELATIONSHIP EVOLUTION\n- Later → State: Must not leak backward.",
            )
        },
        historical_contexts=[
            first_snapshot,
            legacy_later_snapshot,
        ],
    )

    first = await tracker.next_context(
        text="First draft.",
        llm_client=MagicMock(),
        model_name="model",
        target_language="English",
        display_index=1,
        total_chunks=2,
    )
    second = await tracker.next_context(
        text="Second draft.",
        llm_client=MagicMock(),
        model_name="model",
        target_language="English",
        display_index=2,
        total_chunks=2,
    )

    assert "- Alice ↔ Bob: Established romantic couple." in first
    assert "Later → State" not in first
    assert "- Alice ↔ Bob: Established romantic couple." in second
    assert '- Alice → Bob: source form "Bob"' in second
    assert "- Guard → Captain: Temporary scene relationship." in second


@pytest.mark.asyncio
async def test_refinement_snapshot_explicit_delete_removes_durable_relationship():
    first_snapshot = build_novel_context(
        "# GLOBAL LORE",
        (
            "## CURRENT ADDRESSING FORMS\n"
            '- Alice → Bob: source form "Bob" | intimate\n\n'
            "## RELATIONSHIP EVOLUTION\n"
            "- Alice ↔ Bob: Established romantic couple."
        ),
    )
    deleted_snapshot = build_novel_context(
        "# GLOBAL LORE",
        (
            "## CURRENT ADDRESSING FORMS\n"
            "- Alice → Bob: DELETE\n\n"
            "## RELATIONSHIP EVOLUTION\n"
            "- Alice ↔ Bob: DELETE"
        ),
    )
    tracker = RefinementContextTracker(
        prompt_options={
            "novel_context": build_novel_context("# GLOBAL LORE", "")
        },
        historical_contexts=[first_snapshot, deleted_snapshot],
    )

    await tracker.next_context(
        text="First draft.",
        llm_client=MagicMock(),
        model_name="model",
        target_language="English",
        display_index=1,
        total_chunks=2,
    )
    second = await tracker.next_context(
        text="Second draft.",
        llm_client=MagicMock(),
        model_name="model",
        target_language="English",
        display_index=2,
        total_chunks=2,
    )

    assert "Alice → Bob" not in second
    assert "Alice ↔ Bob" not in second


@pytest.mark.asyncio
async def test_standalone_refinement_rebuilds_context_before_refining(monkeypatch):
    from src.core import translator

    events = []

    class FakeClient:
        async def generate(self, prompt, system_prompt=None, **_kwargs):
            if "DRAFT TRANSLATION TO AUDIT" in prompt:
                events.append("reflection")
                return LLMResponse(content=(
                    '<REFLECTION_JSON>{"status":"needs_repair","issues":[{'
                    '"category":"style","severity":"major","source_quote":"",'
                    '"draft_quote":"Mira greeted Rowan.",'
                    '"instruction":"Polish the sentence.",'
                    '"draft_replacement":{"draft":"Mira greeted Rowan.",'
                    '"replacement":"Polished text"},"glossary_update":null,'
                    '"term_replacement":null}]}</REFLECTION_JSON>'
                ))
            if "SENIOR EDITOR CRITIQUE" in prompt:
                events.append("repair")
                return LLMResponse(
                    content="<TRANSLATION>Polished text</TRANSLATION>"
                )
            events.append("context")
            return LLMResponse(
                content=(
                    "[NEW_CHARACTERS]\n"
                    "- Mira: Female protagonist\n\n"
                    "[NEW_GLOSSARY]\n\n"
                    "[DYNAMIC_STATE]\n"
                    "Mira addresses Rowan formally."
                )
            )

        async def detect_thinking_model(self):
            return None

        async def close(self):
            return None

    monkeypatch.setattr(translator, "create_llm_client", lambda *args, **kwargs: FakeClient())

    tracker = RefinementContextTracker(
        prompt_options={
            "auto_update_context": True,
            "novel_context": build_novel_context("# GLOBAL LORE", ""),
        },
        historical_contexts=[None],
    )
    result = await translator.refine_chunks(
        translated_chunks=["Mira greeted Rowan."],
        original_chunks=[{
            "context_before": "",
            "main_content": "Mira greeted Rowan.",
            "context_after": "",
        }],
        target_language="English",
        model_name="model",
        api_endpoint="http://localhost",
        prompt_options=tracker.prompt_options,
        context_tracker=tracker,
    )

    assert result == ["Polished text"]
    assert events == ["context", "context", "reflection"]


@pytest.mark.asyncio
async def test_txt_refine_after_reuses_exact_translation_units(monkeypatch, tmp_path):
    from src.core.refine import txt_refiner

    input_path = tmp_path / "translated.txt"
    output_path = tmp_path / "refined.txt"
    input_path.write_text("One\nTwo", encoding="utf-8")

    context_one = build_novel_context("LORE", "STATE ONE")
    context_two = build_novel_context("LORE", "STATE TWO")
    rows = [
        {
            "chunk_index": 0,
            "status": "completed",
            "translated_text": "One",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(context_one),
            },
        },
        {
            "chunk_index": 1,
            "status": "completed",
            "translated_text": "Two",
            "chunk_data": {
                "context_snapshot": compress_dynamic_state(context_two),
            },
        },
    ]
    checkpoint_manager = MagicMock()
    checkpoint_manager.db.get_chunks.return_value = rows

    monkeypatch.setattr(
        txt_refiner,
        "split_text_into_chunks",
        lambda *args, **kwargs: [{
            "context_before": "",
            "main_content": "One\nTwo",
            "context_after": "",
        }],
    )

    async def fake_refine_chunks(**kwargs):
        assert kwargs["translated_chunks"] == ["One", "Two"]
        assert len(kwargs["original_chunks"]) == 2
        assert kwargs["context_tracker"].historical_contexts == [
            context_one,
            context_two,
        ]
        return ["One refined", "Two refined"]

    monkeypatch.setattr(txt_refiner, "refine_chunks", fake_refine_chunks)

    success = await txt_refiner.refine_txt_file(
        input_filepath=str(input_path),
        output_filepath=str(output_path),
        target_language="English",
        model_name="model",
        checkpoint_manager=checkpoint_manager,
        translation_id="job",
        prompt_options={"novel_context": build_novel_context("LORE", "")},
    )

    assert success is True
    assert output_path.read_text(encoding="utf-8").startswith(
        "One refined\nTwo refined"
    )


@pytest.mark.asyncio
async def test_txt_standalone_refinement_aligns_optional_original(monkeypatch, tmp_path):
    from src.core.refine import txt_refiner

    translated = tmp_path / "translated.txt"
    original = tmp_path / "original.txt"
    output = tmp_path / "refined.txt"
    translated.write_text("Bonjour monde.", encoding="utf-8")
    original.write_text("Hello world.", encoding="utf-8")
    captured = {}

    async def fake_refine_chunks(**kwargs):
        captured.update(kwargs)
        return kwargs["translated_chunks"]

    monkeypatch.setattr(txt_refiner, "refine_chunks", fake_refine_chunks)
    assert await txt_refiner.refine_txt_file(
        input_filepath=str(translated),
        output_filepath=str(output),
        target_language="French",
        refinement_original_path=str(original),
        prompt_options={"source_language": "English"},
    )
    assert captured["original_chunks"][0]["source_content"] == "Hello world."


@pytest.mark.asyncio
async def test_epub_refine_runs_without_stats_callback(monkeypatch, tmp_path):
    from src.core.refine import epub_refiner

    input_path = tmp_path / "input.epub"
    output_path = tmp_path / "output.epub"
    input_path.write_bytes(b"epub")

    client = MagicMock()
    client.close = AsyncMock()
    monkeypatch.setattr(
        epub_refiner,
        "build_refine_client",
        lambda **kwargs: (client, None),
    )
    monkeypatch.setattr(epub_refiner, "_extract_epub", lambda *args, **kwargs: None)
    monkeypatch.setattr(
        epub_refiner,
        "_parse_epub_manifest",
        lambda *args, **kwargs: {"content_files": [], "opf_dir": str(tmp_path)},
    )

    def fake_repackage(*, output_filepath, **kwargs):
        Path(output_filepath).write_bytes(b"refined")

    monkeypatch.setattr(epub_refiner, "_repackage_epub", fake_repackage)

    success = await epub_refiner.refine_epub_file(
        input_filepath=str(input_path),
        output_filepath=str(output_path),
        target_language="English",
        model_name="model",
        stats_callback=None,
    )

    assert success is True
    assert output_path.read_bytes() == b"refined"


@pytest.mark.asyncio
async def test_epub_refine_emits_global_chunk_progress(monkeypatch, tmp_path):
    from src.core.refine import epub_refiner

    input_path = tmp_path / "input.epub"
    output_path = tmp_path / "output.epub"
    input_path.write_bytes(b"epub")
    (tmp_path / "a.xhtml").write_text("<html><body>A</body></html>", encoding="utf-8")
    (tmp_path / "b.xhtml").write_text("<html><body>B</body></html>", encoding="utf-8")

    client = MagicMock()
    client.close = AsyncMock()
    monkeypatch.setattr(
        epub_refiner,
        "build_refine_client",
        lambda **kwargs: (client, None),
    )
    monkeypatch.setattr(epub_refiner, "_extract_epub", lambda *args, **kwargs: None)
    monkeypatch.setattr(
        epub_refiner,
        "_parse_epub_manifest",
        lambda *args, **kwargs: {
            "content_files": ["a.xhtml", "b.xhtml"],
            "opf_dir": str(tmp_path),
        },
    )
    monkeypatch.setattr(
        epub_refiner.etree,
        "parse",
        lambda *args, **kwargs: MagicMock(getroot=lambda: MagicMock()),
    )
    monkeypatch.setattr(
        epub_refiner,
        "_setup_translation",
        lambda *args, **kwargs: ("body", MagicMock(), MagicMock()),
    )
    monkeypatch.setattr(
        epub_refiner,
        "_preserve_tags",
        lambda *args, **kwargs: ("text", {}, ("__PH_", "__")),
    )
    monkeypatch.setattr(
        epub_refiner,
        "_create_chunks",
        lambda *args, **kwargs: [
            {"text": "chunk one", "global_indices": []},
            {"text": "chunk two", "global_indices": []},
        ],
    )

    async def fake_refine_one_xhtml(**kwargs):
        kwargs["stats_callback"]({
            "total_chunks": 2,
            "completed_chunks": 1,
            "failed_chunks": 0,
        })
        kwargs["stats_callback"]({
            "total_chunks": 2,
            "completed_chunks": 2,
            "failed_chunks": 0,
        })
        return True

    monkeypatch.setattr(epub_refiner, "_refine_one_xhtml", fake_refine_one_xhtml)

    def fake_repackage(*, output_filepath, **kwargs):
        Path(output_filepath).write_bytes(b"refined")

    monkeypatch.setattr(epub_refiner, "_repackage_epub", fake_repackage)

    stats_events = []
    success = await epub_refiner.refine_epub_file(
        input_filepath=str(input_path),
        output_filepath=str(output_path),
        target_language="English",
        model_name="model",
        stats_callback=lambda stats: stats_events.append(dict(stats)),
    )

    assert success is True
    assert output_path.read_bytes() == b"refined"
    assert all(event["total_chunks"] == 4 for event in stats_events)
    completed = [event["completed_chunks"] for event in stats_events]
    assert completed == sorted(completed)
    assert completed[:3] == [0, 1, 2]
    assert completed[-1] == 4


def test_epub_refine_after_chapter_mode_uses_spine_units():
    from src.core.refine import epub_refiner

    budget, chapter_mode = epub_refiner._refine_chunking_options(
        {"_refine_after": True, "chapter_mode": True},
        5000,
    )

    assert budget >= 10_000_000
    assert chapter_mode is False


def test_structured_chunk_log_can_hide_internal_token_budget():
    from src.core.epub.xhtml_translator import _create_chunks

    logs = []
    _create_chunks(
        "[id0]Hello[id1]",
        {"[id0]": "<p>", "[id1]": "</p>"},
        10_000_000,
        log_callback=lambda event, message: logs.append((event, message)),
        chunking_note="EPUB spine-file refinement unit(s)",
    )

    assert logs == [
        ("chunks_created", "Created 1 chunks as EPUB spine-file refinement unit(s)")
    ]
    assert "10_000_000" not in logs[0][1]
    assert "10000000" not in logs[0][1]


def test_web_refine_after_uses_one_backend_refinement_phase():
    project_root = Path(__file__).resolve().parents[2]
    batch_controller = (
        project_root / "src" / "web" / "static" / "js"
        / "translation" / "batch-controller.js"
    ).read_text(encoding="utf-8")
    handlers = (
        project_root / "src" / "api" / "handlers.py"
    ).read_text(encoding="utf-8")

    assert "refine: false" in batch_controller
    assert "translation_prompt_options['refine'] = False" in handlers
    assert "refine_success = await refine_file(" in handlers
    assert "refinement_prompt_options['_refine_after'] = True" in handlers


def test_workflow_steps_and_resync_logs_use_the_canonical_ui_channel():
    project_root = Path(__file__).resolve().parents[2]
    tracker = (
        project_root / "src" / "web" / "static" / "js"
        / "translation" / "translation-tracker.js"
    ).read_text(encoding="utf-8")
    generic = (
        project_root / "src" / "core" / "adapters" / "generic_translator.py"
    ).read_text(encoding="utf-8")

    assert "MessageLogger.addStepLog" in tracker
    assert "ui_step === 'context_resync'" in tracker
    assert '"ui_step": "context_resync"' in generic
    assert "emit_update(" in generic


def test_context_resync_save_does_not_replace_latest_with_historical_snapshot():
    project_root = Path(__file__).resolve().parents[2]
    tracker = (
        project_root / "src" / "web" / "static" / "js"
        / "translation" / "translation-tracker.js"
    ).read_text(encoding="utf-8")

    save_block = tracker.split(
        "window.saveContextResync = async function()",
        1,
    )[1]
    assert "NovelContextUI.latestContent = newContent" not in save_block
    assert "NovelContextUI.renderContextTabs(newContent, true)" in save_block
    assert "scope: isGlobal ? 'global_lore' : 'snapshot'" in save_block


def test_chunk_context_snapshot_route_keeps_timeline_snapshot_isolated():
    project_root = Path(__file__).resolve().parents[2]
    routes = (
        project_root / "src" / "api" / "blueprints"
        / "translation_routes.py"
    ).read_text(encoding="utf-8")

    snapshot_block = routes.split(
        "if request.args.get('scope') == 'global_lore':",
        1,
    )[1].split("return jsonify", 1)[0]
    assert "normalize_refinement_context(" in snapshot_block
    assert "plain_text_context = historical_context" in snapshot_block


def test_refinement_context_state_omits_large_ui_payloads():
    project_root = Path(__file__).resolve().parents[2]
    refine_file = (
        project_root / "src" / "core" / "adapters" / "refine_file.py"
    ).read_text(encoding="utf-8")
    tracker = (
        project_root / "src" / "web" / "static" / "js"
        / "translation" / "translation-tracker.js"
    ).read_text(encoding="utf-8")

    assert '"content_omitted": True' in refine_file
    assert '"content_size": len(prompt_options' in refine_file
    context_state_block = tracker.split(
        "data.log_entry.type === 'novel_context_state'",
        1,
    )[1]
    assert "typeof data.log_entry.data.content === 'string'" in context_state_block
    assert "if (hasContextContent)" in context_state_block


def test_running_refinement_can_be_restored_after_browser_refresh():
    project_root = Path(__file__).resolve().parents[2]
    routes = (
        project_root / "src" / "api" / "blueprints" / "translation_routes.py"
    ).read_text(encoding="utf-8")
    tracker = (
        project_root / "src" / "web" / "static" / "js"
        / "translation" / "translation-tracker.js"
    ).read_text(encoding="utf-8")

    assert "'input_filename': (" in routes
    assert "job.input_filename || job.output_filename" in tracker
    assert "fileType: job.file_type || 'txt'" in tracker


def test_max_tokens_setting_is_runtime_reloadable(monkeypatch, tmp_path):
    import src.config as config

    env_path = tmp_path / ".env"
    env_path.write_text("MAX_TOKENS_PER_CHUNK=137\n", encoding="utf-8")
    old_value = config.MAX_TOKENS_PER_CHUNK

    monkeypatch.setattr(config, "_env_file", env_path)
    monkeypatch.setattr(config, "MAX_TOKENS_PER_CHUNK", old_value)
    monkeypatch.setenv("MAX_TOKENS_PER_CHUNK", str(old_value))

    config.reload_config()

    assert config.MAX_TOKENS_PER_CHUNK == 137


def test_web_jobs_use_and_clamp_live_chunk_budget(monkeypatch):
    import src.config as config
    from src.api.blueprints.translation_routes import _clamp_chunk_tokens

    monkeypatch.setattr(config, "MAX_TOKENS_PER_CHUNK", 137)
    monkeypatch.setattr(config, "reload_config", lambda: None)

    assert _clamp_chunk_tokens(None) == 137
    assert _clamp_chunk_tokens(20) == 50
    assert _clamp_chunk_tokens(2000) == 2000
    assert _clamp_chunk_tokens(5000) == 5000


def test_web_jobs_reload_chunk_budget_when_ui_omits_value(monkeypatch):
    import src.config as config
    from src.api.blueprints.translation_routes import _clamp_chunk_tokens

    def reload_config():
        config.MAX_TOKENS_PER_CHUNK = 5000

    monkeypatch.setattr(config, "MAX_TOKENS_PER_CHUNK", 450)
    monkeypatch.setattr(config, "reload_config", reload_config)

    assert _clamp_chunk_tokens(None) == 5000
